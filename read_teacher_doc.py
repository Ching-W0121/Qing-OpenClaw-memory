from docx import Document
import json

# 读取文档
doc = Document('teacher_advice.docx')

# 提取所有内容
content = []
for para in doc.paragraphs:
    if para.text.strip():
        content.append(para.text)

# 提取表格内容
tables = []
for table in doc.tables:
    table_data = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            table_data.append(cells)
    if table_data:
        tables.append(table_data)

# 输出
print("=" * 80)
print("ChatGPT GPT-5.3 老师的建议")
print("=" * 80)

for text in content:
    print(text)

if tables:
    print("\n" + "=" * 80)
    print("表格内容")
    print("=" * 80)
    for i, table in enumerate(tables, 1):
        print(f"\n表格 {i}:")
        for row in table:
            print(" | ".join(row))

print("\n" + "=" * 80)
