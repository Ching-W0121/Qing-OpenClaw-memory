from docx import Document
import os
import sys

# 设置 UTF-8 编码
sys.stdout.reconfigure(encoding='utf-8')

# 获取工作目录
workspace = r"C:\Users\TR\.openclaw\workspace"

# 找到包含"补充"的 docx 文件
for file in os.listdir(workspace):
    if file.endswith('.docx') and '补充' in file:
        filepath = os.path.join(workspace, file)
        print(f"找到文件：{file}")
        print(f"完整路径：{filepath}")
        
        try:
            doc = Document(filepath)
            print(f"\n段落数：{len(doc.paragraphs)}")
            print("\n--- 内容开始 ---\n")
            
            content = []
            for i, p in enumerate(doc.paragraphs):
                if p.text.strip():
                    content.append(f"[{i}] {p.text}")
            
            # 写入到文件（避免编码问题）
            with open('求职 agent 补充 - 提取内容.txt', 'w', encoding='utf-8') as f:
                f.write('\n'.join(content))
            
            print(f"\n内容已保存到：求职 agent 补充 - 提取内容.txt")
            print(f"共 {len(content)} 段")
            
        except Exception as e:
            print(f"读取失败：{e}")
            import traceback
            traceback.print_exc()
