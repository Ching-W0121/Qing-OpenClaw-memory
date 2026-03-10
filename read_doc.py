from docx import Document
import sys

try:
    doc = Document(r"C:\Users\TR\Desktop\青 agent1.4.docx")
    
    print("=" * 60)
    print("ChatGPT 老师的建议")
    print("=" * 60)
    
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            print(f"\n{para.text}")
    
    # 读取表格
    if doc.tables:
        print("\n" + "=" * 60)
        print("表格内容")
        print("=" * 60)
        
        for i, table in enumerate(doc.tables, 1):
            print(f"\n表格 {i}:")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    print(" | ".join(cells))
    
    print("\n" + "=" * 60)
    
except Exception as e:
    print(f"错误：{e}")
    sys.exit(1)
