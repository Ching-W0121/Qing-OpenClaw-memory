from docx import Document
import os

# 获取工作目录
workspace = r"C:\Users\TR\.openclaw\workspace"

# 找到包含"补充"的 docx 文件
for file in os.listdir(workspace):
    if file.endswith('.docx') and '补充' in file:
        filepath = os.path.join(workspace, file)
        print(f"找到文件：{file}")
        print(f"完整路径：{filepath}")
        
        try:
            doc = Document(filepath)
            print(f"\n段落数：{len(doc.paragraphs)}")
            print("\n--- 内容开始 ---\n")
            
            for i, p in enumerate(doc.paragraphs):
                if p.text.strip():
                    print(f"[{i}] {p.text}")
            
            print("\n--- 内容结束 ---\n")
        except Exception as e:
            print(f"读取失败：{e}")
